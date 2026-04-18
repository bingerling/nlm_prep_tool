#!/usr/bin/env python3
"""
nlm_prep.py - NotebookLM Upload Prep Tool

Combines multiple files (md, txt, docx, pdf, svg, html, mp4, wav, mp3) from a directory into a single
consolidated markdown file ready for upload to Google's NotebookLM.

Usage:
    python nlm_prep.py <input_directory> -o <output_file>

Example:
    python nlm_prep.py "G:\\My Study Notes" -o "combined.md"
"""

import os
import sys
import argparse
import xml.etree.ElementTree as ET
import shutil
from pathlib import Path
from typing import List, Tuple, Optional, Dict
from datetime import datetime

# Try to import optional dependencies, provide helpful errors if not installed
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from tinytag import TinyTag
    MEDIA_AVAILABLE = True
except ImportError:
    MEDIA_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False


def parse_args() -> argparse.Namespace:
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(
        description="Combine files into a single document for NotebookLM upload",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python nlm_prep.py "G:\\My Notes" -o combined.md
  python nlm_prep.py ./study-materials -o output.md --verbose
  python nlm_prep.py ./course -o combined.md --media-dir ./media
        """
    )
    parser.add_argument(
        "input_dir",
        help="Directory containing files to combine"
    )
    parser.add_argument(
        "-o", "--output",
        default="nlm_combined.md",
        help="Output filename (default: nlm_combined.md)"
    )
    parser.add_argument(
        "-v", "--verbose",
        action="store_true",
        help="Show detailed progress"
    )
    parser.add_argument(
        "--skip-docx",
        action="store_true",
        help="Skip DOCX files (if python-docx not installed)"
    )
    parser.add_argument(
        "--skip-pdf",
        action="store_true",
        help="Skip PDF files (if PyMuPDF not installed)"
    )
    parser.add_argument(
        "--skip-svg",
        action="store_true",
        help="Skip SVG files"
    )
    parser.add_argument(
        "--skip-media",
        action="store_true",
        help="Skip media files (mp4, wav, mp3)"
    )
    parser.add_argument(
        "--skip-html",
        action="store_true",
        help="Skip HTML files"
    )
    parser.add_argument(
        "--media-dir",
        default="media",
        help="Output directory for media files (default: media)"
    )
    parser.add_argument(
        "--no-media-index",
        action="store_true",
        help="Skip generating media index file"
    )
    
    return parser.parse_args()


def get_supported_files(input_dir: Path) -> List[Tuple[Path, str]]:
    """
    Recursively find all supported files in the directory.
    Returns list of (file_path, relative_path) tuples.
    """
    supported_extensions = {'.md', '.txt', '.docx', '.pdf', '.svg', '.html', '.htm', '.mp4', '.wav', '.mp3'}
    files = []
    
    for file_path in input_dir.rglob('*'):
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            # Get relative path from input directory
            rel_path = file_path.relative_to(input_dir)
            files.append((file_path, str(rel_path)))
    
    # Sort files for consistent ordering
    files.sort(key=lambda x: x[1].lower())
    return files


def extract_docx_content(docx_path: Path) -> str:
    """
    Extract content from DOCX file and convert to Markdown.
    Preserves formatting like headers, bold, italic, lists, and tables.
    """
    if not DOCX_AVAILABLE:
        return f"\n[DOCX file: {docx_path.name} - python-docx not installed, install with: pip install python-docx]\n"
    
    try:
        doc = Document(docx_path)
        markdown_lines = []
        
        for para in doc.paragraphs:
            if not para.text.strip():
                markdown_lines.append("")
                continue
            
            # Determine paragraph style
            style_name = para.style.name if para.style else "Normal"
            text = para.text
            
            # Check for formatting in runs
            formatted_text = ""
            for run in para.runs:
                run_text = run.text
                if run.bold and run.italic:
                    run_text = f"***{run_text}***"
                elif run.bold:
                    run_text = f"**{run_text}**"
                elif run.italic:
                    run_text = f"*{run_text}*"
                formatted_text += run_text
            
            # Use formatted text if available, otherwise plain text
            display_text = formatted_text if formatted_text else text
            
            # Apply heading styles
            if style_name.startswith('Heading 1'):
                markdown_lines.append(f"# {display_text}")
            elif style_name.startswith('Heading 2'):
                markdown_lines.append(f"## {display_text}")
            elif style_name.startswith('Heading 3'):
                markdown_lines.append(f"### {display_text}")
            elif style_name.startswith('Heading 4'):
                markdown_lines.append(f"#### {display_text}")
            elif style_name.startswith('Heading 5'):
                markdown_lines.append(f"##### {display_text}")
            elif style_name.startswith('Heading 6'):
                markdown_lines.append(f"###### {display_text}")
            elif para.style.name == 'List Bullet':
                markdown_lines.append(f"- {display_text}")
            elif para.style.name == 'List Number':
                markdown_lines.append(f"1. {display_text}")
            else:
                markdown_lines.append(display_text)
        
        # Extract tables
        for table in doc.tables:
            markdown_lines.append("\n")
            # Process table rows
            for i, row in enumerate(table.rows):
                row_cells = []
                for cell in row.cells:
                    # Get text from cell, clean it up
                    cell_text = cell.text.strip().replace('\n', ' ')
                    row_cells.append(cell_text)
                
                # Create markdown table row
                markdown_lines.append("| " + " | ".join(row_cells) + " |")
                
                # Add separator after header row
                if i == 0:
                    separator = "|" + "|".join(["---"] * len(row_cells)) + "|"
                    markdown_lines.append(separator)
            
            markdown_lines.append("")
        
        return '\n'.join(markdown_lines)
    
    except Exception as e:
        return f"\n[Error reading DOCX file {docx_path.name}: {str(e)}]\n"


def extract_pdf_content(pdf_path: Path) -> str:
    """
    Extract text content from PDF file.
    Returns text organized by pages.
    """
    if not PDF_AVAILABLE:
        return f"\n[PDF file: {pdf_path.name} - PyMuPDF not installed, install with: pip install pymupdf]\n"
    
    try:
        doc = fitz.open(pdf_path)
        markdown_lines = []
        
        markdown_lines.append(f"*PDF Document: {pdf_path.name}*")
        markdown_lines.append(f"*Pages: {len(doc)}*")
        markdown_lines.append("")
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            
            if text.strip():
                markdown_lines.append(f"### Page {page_num + 1}")
                markdown_lines.append("")
                markdown_lines.append(text)
                markdown_lines.append("")
        
        doc.close()
        return '\n'.join(markdown_lines)
    
    except Exception as e:
        return f"\n[Error reading PDF file {pdf_path.name}: {str(e)}]\n"


def extract_svg_content(svg_path: Path) -> str:
    """
    Extract text content from SVG file.
    Extracts text from <text>, <title>, and <desc> elements.
    """
    try:
        tree = ET.parse(svg_path)
        root = tree.getroot()
        
        markdown_lines = []
        markdown_lines.append(f"*SVG Diagram: {svg_path.name}*")
        markdown_lines.append("")
        
        # Extract title
        title = root.find('.//{http://www.w3.org/2000/svg}title')
        if title is not None and title.text:
            markdown_lines.append(f"**Title:** {title.text}")
            markdown_lines.append("")
        
        # Extract description
        desc = root.find('.//{http://www.w3.org/2000/svg}desc')
        if desc is not None and desc.text:
            markdown_lines.append(f"**Description:** {desc.text}")
            markdown_lines.append("")
        
        # Extract all text elements
        text_elements = root.findall('.//{http://www.w3.org/2000/svg}text')
        if text_elements:
            extracted_texts = []
            for text_elem in text_elements:
                if text_elem.text and text_elem.text.strip():
                    extracted_texts.append(text_elem.text.strip())
            
            if extracted_texts:
                markdown_lines.append("**Text Content:**")
                for text in extracted_texts:
                    markdown_lines.append(f"- {text}")
                markdown_lines.append("")
        
        # If no text found, add a note
        if len(markdown_lines) <= 2:
            markdown_lines.append("*[No extractable text content found in this SVG]*")
        
        return '\n'.join(markdown_lines)
    
    except ET.ParseError as e:
        return f"\n[Error parsing SVG file {svg_path.name}: {str(e)}]\n"
    except Exception as e:
        return f"\n[Error reading SVG file {svg_path.name}: {str(e)}]\n"


def read_text_file(file_path: Path) -> str:
    """Read a text or markdown file."""
    try:
        # Try UTF-8 first
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Fall back to latin-1
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            return f"\n[Error reading file: {str(e)}]\n"


def generate_toc(files: List[Tuple[Path, str]]) -> str:
    """Generate a table of contents from the file list."""
    toc_lines = ["## Table of Contents\n"]
    
    current_dir = ""
    for file_path, rel_path in files:
        # Get directory and filename
        path_parts = Path(rel_path).parts
        dir_name = path_parts[0] if len(path_parts) > 1 else "Root"
        file_name = path_parts[-1]
        
        # Add directory header if changed
        if dir_name != current_dir:
            if current_dir != "":
                toc_lines.append("")
            toc_lines.append(f"### {dir_name}")
            current_dir = dir_name
        
        # Clean filename for display (remove extension)
        display_name = Path(file_name).stem
        # Create anchor link
        anchor = rel_path.replace(' ', '-').replace('\\', '-').replace('/', '-').lower()
        anchor = ''.join(c for c in anchor if c.isalnum() or c == '-')
        
        toc_lines.append(f"- [{display_name}](#{anchor})")
    
    return '\n'.join(toc_lines)


def count_words(text: str) -> int:
    """Count words in text."""
    return len(text.split())


def format_file_size(size_bytes: int) -> str:
    """Format file size in human readable format."""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024.0:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024.0
    return f"{size_bytes:.1f} TB"


def extract_media_metadata(media_path: Path) -> Dict[str, str]:
    """
    Extract metadata from media files (mp4, wav, mp3).
    Returns a dictionary with metadata information.
    """
    if not MEDIA_AVAILABLE:
        return {"error": "tinytag not installed, install with: pip install tinytag"}
    
    try:
        tag = TinyTag.get(media_path)
        metadata = {
            "title": tag.title or "",
            "artist": tag.artist or "",
            "album": tag.album or "",
            "duration": f"{int(tag.duration // 60)}:{int(tag.duration % 60):02d}" if tag.duration else "",
            "bitrate": f"{tag.bitrate} kbps" if tag.bitrate else "",
            "samplerate": f"{tag.samplerate} Hz" if tag.samplerate else "",
            "filesize": format_file_size(tag.filesize) if tag.filesize else "",
        }
        return metadata
    except Exception as e:
        return {"error": str(e)}


def copy_media_file(source: Path, dest_dir: Path, counter: int) -> Tuple[Path, str]:
    """
    Copy a media file to the destination directory with sequential numbering.
    Returns (new_path, new_filename).
    """
    # Create destination directory if it doesn't exist
    dest_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate new filename with sequential number
    new_filename = f"{counter:03d}_{source.name}"
    dest_path = dest_dir / new_filename
    
    # Copy the file
    shutil.copy2(source, dest_path)
    
    return dest_path, new_filename


def extract_html_content(html_path: Path) -> str:
    """
    Extract text content from HTML file and convert to Markdown.
    Preserves headings, paragraphs, lists, and links.
    """
    if not HTML_AVAILABLE:
        return f"\n[HTML file: {html_path.name} - beautifulsoup4 not installed, install with: pip install beautifulsoup4]\n"
    
    try:
        with open(html_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        
        markdown_lines = []
        markdown_lines.append(f"*HTML Document: {html_path.name}*")
        markdown_lines.append("")
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Process elements in order
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li', 'a', 'table', 'pre', 'code', 'blockquote']):
            if element.name == 'h1':
                markdown_lines.append(f"# {element.get_text().strip()}")
            elif element.name == 'h2':
                markdown_lines.append(f"## {element.get_text().strip()}")
            elif element.name == 'h3':
                markdown_lines.append(f"### {element.get_text().strip()}")
            elif element.name == 'h4':
                markdown_lines.append(f"#### {element.get_text().strip()}")
            elif element.name == 'h5':
                markdown_lines.append(f"##### {element.get_text().strip()}")
            elif element.name == 'h6':
                markdown_lines.append(f"###### {element.get_text().strip()}")
            elif element.name == 'p':
                text = element.get_text().strip()
                if text:
                    markdown_lines.append(text)
            elif element.name == 'li':
                # Check if parent is ul or ol
                parent = element.find_parent(['ul', 'ol'])
                if parent and parent.name == 'ol':
                    # For ordered lists, we'll handle numbering separately
                    markdown_lines.append(f"1. {element.get_text().strip()}")
                else:
                    markdown_lines.append(f"- {element.get_text().strip()}")
            elif element.name == 'a':
                href = element.get('href', '')
                text = element.get_text().strip()
                if href and text:
                    markdown_lines.append(f"[{text}]({href})")
            elif element.name == 'pre':
                # Code block
                code_text = element.get_text().strip()
                if code_text:
                    markdown_lines.append(f"```\n{code_text}\n```")
            elif element.name == 'blockquote':
                text = element.get_text().strip()
                if text:
                    markdown_lines.append(f"> {text}")
        
        # Clean up consecutive empty lines
        cleaned_lines = []
        prev_empty = False
        for line in markdown_lines:
            is_empty = not line.strip()
            if is_empty and prev_empty:
                continue
            cleaned_lines.append(line)
            prev_empty = is_empty
        
        return '\n'.join(cleaned_lines)
    
    except Exception as e:
        return f"\n[Error reading HTML file {html_path.name}: {str(e)}]\n"


def generate_media_index(media_files: List[Tuple[Path, str, str, Dict[str, str]]], output_dir: Path) -> None:
    """
    Generate an INDEX.md file for media files.
    media_files is a list of (original_path, new_filename, file_type, metadata) tuples.
    """
    index_lines = []
    index_lines.append("# Media Files Index")
    index_lines.append("")
    index_lines.append("*This file maps original media file locations to their new names in this directory.*")
    index_lines.append("")
    index_lines.append(f"*Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")
    index_lines.append("")
    index_lines.append("---")
    index_lines.append("")
    
    # Summary table
    index_lines.append("## Summary")
    index_lines.append("")
    index_lines.append(f"**Total Media Files:** {len(media_files)}")
    index_lines.append("")
    
    # Count by type
    type_counts = {}
    for _, _, file_type, _ in media_files:
        type_counts[file_type] = type_counts.get(file_type, 0) + 1
    
    if type_counts:
        index_lines.append("**File Types:**")
        for file_type, count in sorted(type_counts.items()):
            index_lines.append(f"- {file_type.upper()}: {count}")
    index_lines.append("")
    index_lines.append("---")
    index_lines.append("")
    
    # Detailed file listing
    index_lines.append("## File Listing")
    index_lines.append("")
    
    for i, (original_path, new_filename, file_type, metadata) in enumerate(media_files, 1):
        index_lines.append(f"### {i}. {new_filename}")
        index_lines.append("")
        index_lines.append(f"**Original Path:** `{original_path}`")
        index_lines.append("")
        index_lines.append(f"**Type:** {file_type.upper()}")
        index_lines.append("")
        
        # Add metadata if available
        if metadata and "error" not in metadata:
            if metadata.get("title"):
                index_lines.append(f"**Title:** {metadata['title']}")
            if metadata.get("artist"):
                index_lines.append(f"**Artist:** {metadata['artist']}")
            if metadata.get("album"):
                index_lines.append(f"**Album:** {metadata['album']}")
            if metadata.get("duration"):
                index_lines.append(f"**Duration:** {metadata['duration']}")
            if metadata.get("bitrate"):
                index_lines.append(f"**Bitrate:** {metadata['bitrate']}")
            if metadata.get("samplerate"):
                index_lines.append(f"**Sample Rate:** {metadata['samplerate']}")
            if metadata.get("filesize"):
                index_lines.append(f"**File Size:** {metadata['filesize']}")
        elif metadata and "error" in metadata:
            index_lines.append(f"**Metadata Error:** {metadata['error']}")
        
        index_lines.append("")
    
    # Write index file
    index_path = output_dir / "INDEX.md"
    with open(index_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(index_lines))


def main():
    args = parse_args()
    
    # Validate input directory
    input_dir = Path(args.input_dir).resolve()
    if not input_dir.exists():
        print(f"Error: Directory not found: {input_dir}")
        sys.exit(1)
    if not input_dir.is_dir():
        print(f"Error: Not a directory: {input_dir}")
        sys.exit(1)
    
    # Check for optional dependencies
    if not DOCX_AVAILABLE and not args.skip_docx:
        print("Warning: python-docx not installed. DOCX files will be skipped.")
        print("To install: pip install python-docx")
        print("Or use --skip-docx to suppress this warning")
        print()
    
    if not PDF_AVAILABLE and not args.skip_pdf:
        print("Warning: PyMuPDF not installed. PDF files will be skipped.")
        print("To install: pip install pymupdf")
        print("Or use --skip-pdf to suppress this warning")
        print()
    
    if not MEDIA_AVAILABLE and not args.skip_media:
        print("Warning: tinytag not installed. Media files will be skipped.")
        print("To install: pip install tinytag")
        print("Or use --skip-media to suppress this warning")
        print()
    
    if not HTML_AVAILABLE and not args.skip_html:
        print("Warning: beautifulsoup4 not installed. HTML files will be skipped.")
        print("To install: pip install beautifulsoup4")
        print("Or use --skip-html to suppress this warning")
        print()
    
    # Find all supported files
    if args.verbose:
        print(f"Scanning: {input_dir}")
    
    files = get_supported_files(input_dir)
    
    if not files:
        print(f"No supported files found in {input_dir}")
        print("Supported formats: .md, .txt, .docx, .pdf, .svg, .html, .htm, .mp4, .wav, .mp3")
        sys.exit(1)
    
    # Group files by type for reporting
    md_files = [f for f in files if f[0].suffix.lower() == '.md']
    txt_files = [f for f in files if f[0].suffix.lower() == '.txt']
    docx_files = [f for f in files if f[0].suffix.lower() == '.docx']
    pdf_files = [f for f in files if f[0].suffix.lower() == '.pdf']
    svg_files = [f for f in files if f[0].suffix.lower() == '.svg']
    html_files = [f for f in files if f[0].suffix.lower() in ('.html', '.htm')]
    media_files = [f for f in files if f[0].suffix.lower() in ('.mp4', '.wav', '.mp3')]
    
    print(f"Found {len(files)} files:")
    print(f"  - Markdown: {len(md_files)}")
    print(f"  - Text: {len(txt_files)}")
    print(f"  - DOCX: {len(docx_files)}")
    print(f"  - PDF: {len(pdf_files)}")
    print(f"  - SVG: {len(svg_files)}")
    print(f"  - HTML: {len(html_files)}")
    print(f"  - Media: {len(media_files)}")
    print()
    
    # Setup media directory
    media_dir = Path(args.media_dir).resolve()
    media_counter = 0
    media_file_records = []  # For index generation
    
    # Generate output
    output_lines = []
    output_lines.append(f"# NotebookLM Combined Notes")
    output_lines.append(f"")
    output_lines.append(f"*Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")
    output_lines.append(f"*Source: {input_dir}*")
    output_lines.append(f"")
    
    # Add table of contents
    toc = generate_toc(files)
    output_lines.append(toc)
    output_lines.append("")
    output_lines.append("---")
    output_lines.append("")
    
    # Process each file
    total_words = 0
    processed_count = 0
    skipped_count = 0
    
    for file_path, rel_path in files:
        if args.verbose:
            print(f"Processing: {rel_path}")
        
        # Create section header with anchor
        anchor = rel_path.replace(' ', '-').replace('\\', '-').replace('/', '-').lower()
        anchor = ''.join(c for c in anchor if c.isalnum() or c == '-')
        
        output_lines.append(f'<a name="{anchor}"></a>')
        output_lines.append(f"## {rel_path}")
        output_lines.append("")
        
        # Read content based on file type
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.docx':
            if args.skip_docx:
                content = f"\n[DOCX file skipped: {file_path.name}]\n"
                skipped_count += 1
            else:
                content = extract_docx_content(file_path)
                processed_count += 1
        elif file_ext == '.pdf':
            if args.skip_pdf:
                content = f"\n[PDF file skipped: {file_path.name}]\n"
                skipped_count += 1
            else:
                content = extract_pdf_content(file_path)
                processed_count += 1
        elif file_ext == '.svg':
            if args.skip_svg:
                content = f"\n[SVG file skipped: {file_path.name}]\n"
                skipped_count += 1
            else:
                content = extract_svg_content(file_path)
                processed_count += 1
        elif file_ext in ('.html', '.htm'):
            if args.skip_html:
                content = f"\n[HTML file skipped: {file_path.name}]\n"
                skipped_count += 1
            else:
                content = extract_html_content(file_path)
                processed_count += 1
        elif file_ext in ('.mp4', '.wav', '.mp3'):
            if args.skip_media:
                content = f"\n[Media file skipped: {file_path.name}]\n"
                skipped_count += 1
            else:
                # Copy media file and generate metadata
                media_counter += 1
                _, new_filename = copy_media_file(file_path, media_dir, media_counter)
                metadata = extract_media_metadata(file_path)
                
                # Record for index
                media_file_records.append((rel_path, new_filename, file_ext[1:], metadata))
                
                # Generate markdown content for media file
                content_lines = []
                content_lines.append(f"*Media File: {file_path.name}*")
                content_lines.append("")
                content_lines.append(f"**Type:** {file_ext[1:].upper()}")
                content_lines.append(f"**Copied to:** `{media_dir.name}/{new_filename}`")
                
                if metadata and "error" not in metadata:
                    if metadata.get("title"):
                        content_lines.append(f"**Title:** {metadata['title']}")
                    if metadata.get("artist"):
                        content_lines.append(f"**Artist:** {metadata['artist']}")
                    if metadata.get("album"):
                        content_lines.append(f"**Album:** {metadata['album']}")
                    if metadata.get("duration"):
                        content_lines.append(f"**Duration:** {metadata['duration']}")
                    if metadata.get("bitrate"):
                        content_lines.append(f"**Bitrate:** {metadata['bitrate']}")
                    if metadata.get("samplerate"):
                        content_lines.append(f"**Sample Rate:** {metadata['samplerate']}")
                    if metadata.get("filesize"):
                        content_lines.append(f"**File Size:** {metadata['filesize']}")
                elif metadata and "error" in metadata:
                    content_lines.append(f"**Metadata Error:** {metadata['error']}")
                
                content = '\n'.join(content_lines)
                processed_count += 1
        else:
            content = read_text_file(file_path)
            processed_count += 1
        
        output_lines.append(content)
        output_lines.append("")
        output_lines.append("---")
        output_lines.append("")
        
        total_words += count_words(content)
    
    # Join all lines
    final_content = '\n'.join(output_lines)
    
    # Write output file
    output_path = Path(args.output).resolve()
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(final_content)
    except Exception as e:
        print(f"Error writing output file: {e}")
        sys.exit(1)
    
    # Generate media index if media files were processed
    if media_file_records and not args.no_media_index:
        generate_media_index(media_file_records, media_dir)
        print(f"Media index generated: {media_dir / 'INDEX.md'}")
    
    # Report results
    file_size = output_path.stat().st_size
    
    print("=" * 50)
    print("Processing Complete!")
    print("=" * 50)
    print(f"Output file: {output_path}")
    print(f"File size: {format_file_size(file_size)}")
    print(f"Total words: {total_words:,}")
    print(f"Files processed: {processed_count}")
    if skipped_count > 0:
        print(f"Files skipped: {skipped_count}")
    if media_counter > 0:
        print(f"Media files copied: {media_counter} to {media_dir}")
    print()
    print("Your file is ready to upload to NotebookLM!")
    print("Simply upload this single file as a source in your notebook.")


if __name__ == "__main__":
    main()
